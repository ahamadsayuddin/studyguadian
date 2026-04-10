import base64
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_collaboration_doc():
    # 1. Mermaid Code for Collaboration
    mermaid_code = """graph LR
    U((User))
    V[Views]
    S[Services]
    M[(Models/DB)]
    AI{AI Engine}
    G{Guardian API}

    U -- "1. Request" --> V
    V -- "2. Auth" --> M
    V -- "3. Logic" --> S
    S -- "4. AI" --> AI
    S -- "5. Log" --> M
    S -- "6. Alert" --> G
    V -- "7. Response" --> U"""

    # 2. Get Diagram Image from Mermaid.ink
    print("Generating collaboration diagram image...")
    graphbytes = mermaid_code.encode("utf8")
    base64_bytes = base64.b64encode(graphbytes)
    base64_string = base64_bytes.decode("ascii")
    url = "https://mermaid.ink/img/" + base64_string
    
    try:
        response = requests.get(url)
        if response.status_code == 200:
            with open("temp_collaboration.png", "wb") as f:
                f.write(response.content)
            image_path = "temp_collaboration.png"
        else:
            print(f"Failed to fetch image: {response.status_code}")
            image_path = None
    except Exception as e:
        print(f"Error fetching image: {e}")
        image_path = None

    # 3. Create Word Doc
    print("Creating Word document...")
    doc = Document()
    
    # Title
    title = doc.add_heading('Collaboration Diagram - Study Assistant', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_paragraph("This diagram illustrates the structural relationships and message exchanges between the core components of the Study Assistant.")

    # Diagram
    if image_path:
        doc.add_heading('Collaboration / Communication Model', level=1)
        doc.add_picture(image_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Description
    doc.add_heading('Key Collaborations', level=1)
    
    sections = [
        ("Views & Services", "Views delegate complex logic, such as AI analysis and messaging, to the Services layer to ensure clean separation of concerns."),
        ("Services & Models", "Services collaborate with Models to persist interaction logs, outbound message statuses, and performance data."),
        ("Internal & External Engines", "The system collaborates with AI (Gemini/Groq) for content generation and Twilio for real-time guardian notifications.")
    ]
    
    for head, text in sections:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{head}: ")
        run.bold = True
        p.add_run(text)

    # Save
    doc_name = "Collaboration_Diagram.docx"
    doc.save(doc_name)
    print(f"Saved {doc_name}")
    
    # Cleanup
    if image_path and os.path.exists(image_path):
        os.remove(image_path)

if __name__ == "__main__":
    generate_collaboration_doc()
